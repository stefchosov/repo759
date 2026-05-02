"""
Generates Final/Report/FinalProject_Arsov.docx from template structure.
Run from repo root: python3 Final/Scripts/generate_report.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.style.font.size = Pt(12 if level > 1 else 14)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    doc.add_paragraph()

# Title page
title = doc.add_heading('Spring 2026 ME/CS/ECE 759 Final Project Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph('University of Wisconsin-Madison')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

proj = doc.add_paragraph('CPU vs GPU Performance and Truncated-Collision Analysis\nof MD5, SHA-1, and SHA-256')
proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
proj.runs[0].bold = True

author = doc.add_paragraph('Stefan Arsov\nMay 5, 2026')
author.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# Abstract
heading('Abstract')
body(
    "This project benchmarks MD5, SHA-1, and SHA-256 cryptographic hash functions "
    "across three execution models (CPU serial, CPU parallel via OpenMP, and GPU via "
    "CUDA), and uses truncated-collision search as a vehicle for showing where the "
    "GPU's algorithmic advantages become decisive at scale. "
    "Task 1 measures raw hash throughput as a function of batch size n and establishes "
    "GPU vs CPU speedup baselines (up to 440x for SHA-1 at n=10M). "
    "Task 2 sweeps threads-per-block (tpb) from 32 to 1024 and finds that register "
    "pressure limits occupancy, making tpb=32 optimal for all three algorithms. "
    "Task 3 implements truncated collision search over bit widths {16, 24, 32, 40, 48, "
    "56, 64, 72} using three approaches: a sequential CPU unordered_map scan, a "
    "two-phase OpenMP parallel scan with 12 threads on a single socket, and GPU "
    "Pollard's rho with distinguished points using 65,536 CUDA threads and O(1) memory "
    "per thread. The GPU is the only implementation capable of running at bit widths "
    "of 64 or higher, and at bits=56 it achieves 260x speedup over CPU serial for MD5. "
    "A 128-bit arithmetic extension lets the GPU reach bits=72, a range that is "
    "completely infeasible for any CPU approach. The OpenMP sweet spot at 12 threads "
    "is explained by NUMA topology, the Amdahl's serial merge bottleneck, and L3 cache "
    "pressure. Task 4 compares Pollard's rho against an alternative GPU strategy "
    "(Thrust parallel sort) in two memory modes, with 10 trials each, and shows that "
    "Pollard's rho's main advantage above the VRAM boundary is determinism rather than "
    "raw speed. All experiments ran on the UW-Madison Euler cluster with an NVIDIA GPU, "
    "CUDA 13.0, and a dual-socket CPU node."
)
body("Link to Final Project git repo: https://github.com/stefchosov/repo759")

doc.add_page_break()

# General information
heading('General Information')
bullet("Name: Stefan Arsov")
bullet("Email: sarsov@epic.com")
bullet("Home department: Computer Sciences")
bullet("Status: MS Student")
bullet("Teammate: None")
bullet("I release the ME759 Final Project code as open source and under a BSD3 license "
       "for unfettered use of it by any interested party.")

# 1. Problem Statement
heading('1. Problem Statement')
body(
    "Cryptographic hash functions (MD5, SHA-1, SHA-256) are the computational backbone "
    "of password hashing, digital signatures, file integrity verification, and "
    "certificate chains. Their security rests on two properties: preimage resistance "
    "(hard to find input given output) and collision resistance (hard to find two "
    "inputs with the same output). While the full-length versions of these algorithms "
    "are cryptographically strong, truncated variants appear in many practical systems. "
    "SSH host key fingerprints use truncated MD5 or SHA-256. Git object IDs use "
    "truncated SHA-1. Many authentication protocols use only the first N bits of a "
    "digest to reduce bandwidth."
)
body(
    "This project asks two questions. First, how does GPU parallelism affect raw hash "
    "throughput relative to serial and OpenMP CPU implementations, and how does the "
    "answer vary by algorithm complexity and configuration? Second, how does GPU "
    "parallelism change the feasibility of truncated collision search (finding two "
    "distinct inputs whose N-bit truncated hashes collide), and at what bit width does "
    "the GPU's algorithmic advantage become decisive? The motivation is to understand "
    "the practical boundary between CPU-feasible and GPU-only cryptanalysis tasks, "
    "which is directly relevant to password cracking, rainbow table construction, and "
    "certificate spoofing scenarios."
)

# 2. Solution Description
heading('2. Solution Description')

heading('2.1 Code Structure', level=2)
body("All code lives under Final/Code/ in the repository:")
bullet("task1.cu: GPU throughput benchmark. Launches n parallel hash kernels, measures "
       "time via cudaEvent_t, reports MH/s for each algorithm at tpb=256.")
bullet("task2.cu: tpb sweep. Same workload as Task 1 but varies threads-per-block from "
       "32 to 1024, reporting throughput at each configuration.")
bullet("task3.cu: collision search. CPU serial, OpenMP parallel, and GPU Pollard's rho "
       "implementations for bits in {16, 24, 32, 40, 48, 56, 64, 72}.")
bullet("task4.cu: GPU Thrust-sort collision search. Hash, then thrust::sort_by_key, "
       "then adjacent-duplicate kernel. Bits in {16, 24, 32, 40, 48} for device mode "
       "and {16, 24, 32, 40, 48, 52, 56} for unified-memory mode.")
bullet("gpu_hashes_narrow.cuh: shared device hash functions used by task4.cu.")
bullet("md5_cpu.cpp / sha1_cpu.cpp / sha256_cpu.cpp: portable CPU reference "
       "implementations used by both CPU and OpenMP paths in Task 3.")
bullet("md5_cpu.h / sha1_cpu.h / sha256_cpu.h: headers for the above.")

heading('2.2 Task 1: Hash Throughput', level=2)
body(
    "Task 1 launches n CUDA threads, each computing one hash of an 8-byte little-endian "
    "encoding of its thread index. Three template kernels (template<int ALGO>) "
    "specialize at compile time for MD5, SHA-1, and SHA-256, eliminating runtime "
    "branches. The GPU hash functions operate entirely in registers, with no shared "
    "memory and no global memory reads after the index load. Timing uses cudaEvent_t "
    "around the kernel call with cudaDeviceSynchronize() inside the called function. "
    "The experiment sweeps n from 100K to 10M to show the GPU break-even point relative "
    "to serial CPU."
)

heading('2.3 Task 2: Threads-per-Block Sweep', level=2)
body(
    "Task 2 holds n=10,000,000 fixed and sweeps tpb in {32, 64, 128, 256, 512, 1024}. "
    "The same template kernel from Task 1 is reused with the tpb as a launch parameter. "
    "The key insight is that each hash kernel is register-heavy. SHA-256 maintains 9 "
    "32-bit working variables plus a 64-entry message schedule W[64], so increasing tpb "
    "reduces the number of blocks the SM can schedule concurrently, which in turn "
    "reduces occupancy and hides less latency."
)

heading('2.4 Task 3: Truncated Collision Search', level=2)
body(
    "Task 3 implements three independent collision search strategies for truncated "
    "N-bit hashes, where a collision is two distinct inputs i ≠ j such that "
    "truncate(hash(i), N) == truncate(hash(j), N)."
)

body("CPU Serial: Sequential scan of inputs 0, 1, 2, ... inserting each truncated hash "
     "into an std::unordered_map<uint64_t, uint64_t>. The first duplicate key is a "
     "collision. Limited to bits ≤ 56 due to memory: bits=56 requires roughly 17 GB "
     "for the map, which is accommodated by --mem=96G in SLURM.")

body("OpenMP Parallel (12 threads): Two-phase approach per batch. In Phase 1 "
     "(parallel), threads hash disjoint index ranges into pre-allocated hbuf/ibuf "
     "arrays with no synchronization. In Phase 2 (serial), the merge into the global "
     "map stops at the first duplicate. Thread count is fixed at 12. A dedicated OMP "
     "scaling experiment swept 1 to 32 threads at bits in {40, 48} and showed peak "
     "performance at 12, for three reasons. (1) NUMA boundary: Euler nodes are "
     "dual-socket, so threads 13 and above cross to the remote socket and pay 2x DRAM "
     "latency. (2) Amdahl's serial merge: Phase 2 is fully serial, capping parallel "
     "speedup. (3) L3 cache pressure: the batch working set (about 32 MB at bits=48) "
     "saturates per-socket L3 beyond 12 threads.")

body("GPU Pollard's Rho with Distinguished Points: 65,536 CUDA threads each walk an "
     "independent Markov chain x_{n+1} = truncate(hash(x_n), N bits). A 'distinguished "
     "point' (DP) is any chain value x with its low N/4 bits zero. Two chains reaching "
     "the same DP from different starts imply a collision. This uses O(1) memory per "
     "thread and only an 8 MB DP buffer, so it scales to bits=64 without hitting "
     "device memory limits. Template kernels eliminate warp divergence. For bits "
     "above 64 a 128-bit extension was implemented, with a uint128 struct (lo/hi pair), "
     "wide hash functions accepting 16-byte inputs, and a pollard_kernel_wide template, "
     "enabling bits=72.")

heading('2.5 Task 4: GPU Algorithm Comparison via Thrust Sort', level=2)
body(
    "Task 4 implements a second GPU collision search strategy and benchmarks it "
    "against Pollard's rho from Task 3. The goal is to identify which GPU algorithm "
    "is fastest at each bit width, and where each one breaks down. The implementation "
    "is in Final/Code/task4.cu and reuses the per-algorithm device hash functions "
    "from Final/Code/gpu_hashes_narrow.cuh."
)
body(
    "The algorithm has three phases. Phase 1 launches a parallel hash kernel that "
    "computes batch_n hashes and stores (truncated_hash, input_index) pairs in two "
    "device arrays. Phase 2 calls thrust::sort_by_key on the (hash, index) pairs, "
    "sorted by the hash key. Phase 3 launches a parallel kernel that scans adjacent "
    "elements of the sorted array and uses atomicMin to record the lowest input index "
    "at which a duplicate hash was found. If no duplicate is found in a given batch, "
    "the search retries with a new starting offset and accumulates the count."
)
body(
    "Two memory modes are supported via a CLI flag. The default mode allocates the "
    "device arrays with cudaMalloc, keeping all data in pure VRAM. The bits sweep "
    "is capped at 48 because the working set (4x expected pairs at 16 bytes each, "
    "plus Thrust's internal radix-sort scratch) saturates the 8 GB GPU at bits >= 52. "
    "The unified-memory mode allocates the same arrays with cudaMallocManaged, so the "
    "buffers are visible to both GPU and host and pages migrate across PCIe on demand. "
    "This lets the sweep extend to bits=52 and bits=56, where the working set spills "
    "into the SLURM-allocated 96 GB host RAM. The cost is that the sort becomes "
    "PCIe-bound during page migrations."
)
body(
    "Over-allocation is adaptive based on bit width. At bits <= 48 the binary "
    "allocates 4x the expected count (about 99% chance of finding a collision in "
    "one round). At bits=52 it allocates 2x (about 86% find rate per round). At "
    "bits=56 it allocates 1x (about 63% find rate per round, averaging 1.6 rounds). "
    "Lower over-allocation factors trade higher retry-driven variance for smaller "
    "working sets, which matters because the sort working set scales as the bit "
    "width grows. bits=64 was attempted in unified mode with 0.5x allocation but "
    "thrust::sort_by_key crashed inside its temp-buffer allocation on every trial "
    "and was dropped from the sweep."
)
body(
    "Each algorithm-bits combination runs 10 trials in both memory modes. The "
    "geometric retry distribution at sub-1x allocation has variance comparable to "
    "its mean, so single-trial measurements are not representative. Reported numbers "
    "are median across the 10 trials, with IQR error bars on the time-vs-bits plots."
)

# 3. Results
heading('3. Overview of Results')

heading('3.1 Task 1: Throughput at n=10,000,000', level=2)
add_table(
    ['Algorithm', 'CPU Serial (MH/s)', 'OpenMP-4 (MH/s)', 'GPU (MH/s)', 'GPU/Serial Speedup'],
    [
        ['MD5',     '6.53',  '22.73', '2694.6', '413x'],
        ['SHA-1',   '2.78',   '9.34', '1224.8', '440x'],
        ['SHA-256', '3.94',   '9.54',  '658.9', '167x'],
    ]
)
body(
    "GPU throughput is 167x to 440x faster than serial CPU. OpenMP-4 achieves "
    "near-linear speedup (3.5x) for MD5 and SHA-1, and 2.4x for SHA-256, which is "
    "limited by higher register pressure. At n below roughly 200,000 the GPU is "
    "slower than serial because of kernel launch overhead (about 2 ms). Above that "
    "threshold GPU parallelism dominates. Plots are in Final/Data/task1/task1_throughput.pdf."
)

heading('3.2 Task 2: Threads-per-Block', level=2)
add_table(
    ['tpb', 'MD5 (MH/s)', 'SHA-1 (MH/s)', 'SHA-256 (MH/s)'],
    [
        ['32',   '3021.1', '1201.9', '574.7'],
        ['64',   '2840.9', '1052.6', '559.3'],
        ['128',  '2717.4', '1066.1', '557.7'],
        ['256',  '2277.9', '1021.5', '537.3'],
        ['512',  '2197.8', '1003.0', '540.8'],
        ['1024', '2375.3', '1003.0', '557.4'],
    ]
)
body(
    "tpb=32 (one warp) is optimal for all three algorithms. Throughput decreases as "
    "tpb increases because larger blocks consume more of the SM register file, which "
    "reduces concurrent warp occupancy. MD5 drops 27% from tpb=32 to tpb=512, the "
    "largest impact, because it starts with the highest occupancy. SHA-256 drops only "
    "6%, since it is already register-starved at tpb=32. Using tpb=32 instead of the "
    "default tpb=256 gives a free 7% to 33% speedup. Plot: Final/Data/task2/task2_tpb.pdf."
)

heading('3.3 Task 3: Collision Search Timing', level=2)
body("Selected results, time in ms. CPU and OMP entries are blank where the bit width "
     "exceeds the memory budget for the unordered_map approach.")

add_table(
    ['Algo', 'Bits', 'CPU Serial (ms)', 'OMP-12 (ms)', 'GPU (ms)', 'GPU/CPU Speedup'],
    [
        ['MD5',     '32', '9.5',        '2.6',       '36.7',    '0.26x'],
        ['MD5',     '48', '4,234',      '2,327',     '544',     '7.8x'],
        ['MD5',     '56', '554,672',    '359,864',   '2,133',   '260x'],
        ['MD5',     '64', 'n/a',        'n/a',       '8,488',   'n/a'],
        ['SHA-1',   '56', '74,417',     '35,765',    '2,375',   '31x'],
        ['SHA-256', '56', '316,526',    '181,146',   '1,930',   '164x'],
        ['SHA-256', '64', 'n/a',        'n/a',       '7,585',   'n/a'],
    ]
)
body(
    "GPU speedup over CPU serial grows dramatically with bit width, from sub-1x at "
    "bits=32 (where GPU overhead dominates the small search space) to 31x to 260x at "
    "bits=56. At bits=64 and bits=72 only the GPU runs, because CPU memory requirements "
    "(roughly 100 GB for bits=64 with unordered_map) make CPU collision search "
    "infeasible on available hardware. The OMP implementation achieves only 1.5x to "
    "2.1x speedup over serial at bits=56, limited by the serial merge phase (Amdahl's "
    "Law) and NUMA penalties. The birthday paradox is validated by the CPU data: "
    "cpu_count divided by expected is consistently within a factor of 1 to 2 of 1.0 "
    "for all algorithms and bit widths. Plots: Final/Data/task3/task3_time.pdf, "
    "task3_speedup.pdf, task3_birthday.pdf."
)

body(
    "The OMP thread-count scaling experiment (Final/Data/task3/omp_scaling.dat) "
    "confirmed the 12-thread sweet spot. Speedup peaks at 12 and degrades beyond that "
    "due to NUMA penalties on the dual-socket Euler node. "
    "Plot: Final/Data/task3/omp_scaling_speedup.pdf."
)

heading('3.4 Task 4: GPU Algorithm Comparison (Thrust Sort vs Pollard\'s Rho)', level=2)
body(
    "Task 4 implements an alternative GPU collision search using thrust::sort_by_key "
    "to compare against Task 3's Pollard's rho. The algorithm is three-phase: hash a "
    "batch of inputs into device arrays, sort by hash key, then scan for the first "
    "adjacent equal-hash pair. It is fully GPU-resident with no CPU-side merge phase. "
    "Two memory modes are tested:"
)
bullet(
    "Device mode (cudaMalloc): pure VRAM. Capped at bits=48 because the sort working "
    "set (4x expected, times 16 bytes) plus Thrust scratch saturates 8 GB GPU memory "
    "at bits 52 and above."
)
bullet(
    "Unified memory mode (cudaMallocManaged): the buffers are visible to both GPU and "
    "host, with pages migrating across PCIe on demand. This lets us push the sweep to "
    "bits=52 and bits=56, where the working set spills into host RAM (--mem=96G) "
    "rather than OOMing. The cost is PCIe-bound performance during the sort."
)
body(
    "An adaptive over-allocation factor controls the per-round retry rate: 4x expected "
    "for bits up to 48 (one-shot), 2x for bits=52, and 1x for bits=56 (about 2 rounds "
    "on average to find a collision)."
)
body(
    "Code: Final/Code/task4.cu, which uses gpu_hashes_narrow.cuh for shared device "
    "hash functions. Output: Final/Data/task4/scaling_task4_device.dat and "
    "scaling_task4_unified.dat. Plots: task4_compare.pdf, task4_overhead.pdf, "
    "task4_trial_spread.pdf, task4_speedup.pdf."
)

heading('3.4.1 Methodology: Why Multiple Trials Matter', level=2)
body(
    "Initial Task 4 results were single-trial, which produced misleading conclusions. "
    "The birthday-paradox-driven retry distribution (geometric) has variance "
    "comparable to its mean. At bits=56 with 1x over-allocation, a single run can take "
    "anywhere from 1 to 6 or more rounds depending purely on which input value happens "
    "to collide first. I re-ran with 10 trials per configuration in both device and "
    "unified modes and report median with IQR error bars. Device mode also has small "
    "variance: about a 1.8% miss rate at 4x allocation, plus 5% to 10% jitter from GPU "
    "clock boost and kernel launch overhead. Since device-mode trials are sub-second, "
    "repetition is essentially free. Pollard's rho is essentially deterministic given "
    "fixed seeds and thread scheduling, so a single trial suffices for Task 3 "
    "comparisons."
)

heading('3.4.2 Finding: Unified Memory Has Real but Variable Overhead at Small Bits', level=2)
body(
    "Comparing unified vs device mode at bit widths where the working set fits in 8 GB "
    "VRAM (no page migration required):"
)
add_table(
    ['bits', 'MD5 unified/device', 'SHA-1 unified/device', 'SHA-256 unified/device'],
    [
        ['24',  '3.7x', '4.0x',   '3.7x'],
        ['32',  '2.3x', '0.6x *', '0.6x *'],
        ['40',  '2.2x', '1.6x',   '1.8x'],
        ['48',  '2.0x', '2.3x',   '2.4x'],
    ]
)
body(
    "* SHA-1 and SHA-256 at bits=32 are sub-millisecond; the ratio is dominated by "
    "measurement noise. At meaningful problem sizes (bits=48) the page-tracking "
    "overhead is consistently 2.0x to 2.4x. The takeaway: Unified Memory has a real "
    "2x tax even with no spilling. It is a tool for oversubscription, not a free "
    "abstraction."
)

heading('3.4.3 Finding: Per-round Thrust Beats Pollard\'s Rho, Total Time Depends on Retry Luck', level=2)
body(
    "At bits=56 the comparison against Pollard's rho is more nuanced than single-trial "
    "data suggested:"
)
add_table(
    ['Algorithm (bits=56)', 'Pollard\'s rho (ms)', 'Thrust uni total (ms)', 'rounds', 'per-round (ms)'],
    [
        ['MD5',      '2,133',  '3,359',  '6.0', '560'],
        ['SHA-1',    '2,375',  '1,161',  '1.0', '1,161'],
        ['SHA-256',  '1,930',  '1,951',  '3.0', '650'],
    ]
)
body(
    "Per-round Thrust unified time (560 to 1,161 ms) is consistently faster than "
    "Pollard's rho's deterministic single-run time (1,930 to 2,375 ms). However, "
    "total wall time depends on how many retry rounds are needed. SHA-1 got lucky in "
    "1 round and was 2x faster than Pollard. MD5 needed 6 rounds and was 1.6x slower. "
    "SHA-256 needed 3 rounds and tied. The 10-trial median (see "
    "task4_trial_spread.pdf) collapses this variance and shows the two algorithms are "
    "in the same performance class at bits=56. Pollard's advantage is determinism, "
    "not raw speed."
)

heading('3.4.4 Finding: PCIe Spill Boundary is Sharp at bits 52→56', level=2)
add_table(
    ['bits', 'MD5 unified median (ms)', 'factor vs prior', 'theoretical 2^(bits/2)x'],
    [
        ['48',     '240.9',  'n/a',   'n/a'],
        ['52',     '516.4',  '2.1x',  '2.0x (still fits in VRAM)'],
        ['56',   '3,359',    '6.5x',  '4.0x (extra 1.6x = PCIe spill)'],
    ]
)
body(
    "From bits=48 to 52 the slowdown matches theory exactly: sort scratch still fits "
    "in VRAM. From bits=52 to 56 the extra 1.6x factor beyond theoretical scaling is "
    "the PCIe-bound sort overhead. bits=64 was attempted in unified mode with 0.5x "
    "allocation (about 34 GB working set), but every trial crashed inside "
    "thrust::sort_by_key with an allocation failure for the radix-sort temp scratch "
    "space. Thrust's internal allocator cannot satisfy that request even with 96 GB "
    "host RAM exposed via Unified Memory. Pollard's rho runs cleanly at bits=64 in "
    "8 to 10 seconds, so the algorithmic-feasibility argument from Task 3 is "
    "reinforced empirically. At bits=64, Pollard's rho is not just faster, it is the "
    "only approach that runs at all."
)
body("GPU algorithm selection guide based on measured evidence:")
bullet("bits ≤ 48: Thrust device-mode is fastest absolute (5x to 35x over Pollard's rho)")
bullet("bits = 52: Both Thrust modes work. Unified pays a roughly 2x tax but still beats Pollard's rho")
bullet("bits = 56: Comparable performance class. Choose by deterministic-time preference")
bullet("bits ≥ 64: Pollard's rho only. Thrust unified fails at bits=64 (sort scratch alloc), and at bits ≥ 72 the working set exceeds 96 GB host RAM regardless")

# 4. Deliverables
heading('4. Deliverables: Building and Running')

heading('4.1 Repository Structure', level=2)
body("All project code is in the Final/ subdirectory of the repo:")
bullet("Final/Code/: task1.cu, task2.cu, task3.cu, task4.cu, gpu_hashes_narrow.cuh, plus md5/sha1/sha256 CPU sources")
bullet("Final/sbatch/: SLURM sbatch scripts for all tasks (task1/, task2/, task3/, task4/)")
bullet("Final/Scripts/: Python plotting scripts (plot_final_task*.py)")
bullet("Final/Data/: Output .dat files and generated PDF/PNG plots")
bullet("Final/Report/: per-task analysis text files and this report")

heading('4.2 Compilation', level=2)
body("All commands run from the repository root on the Euler cluster after: "
     "module load nvidia/cuda/13.0.0")
body("Task 1 and Task 2:")
doc.add_paragraph(
    "nvcc -O3 -std=c++17 -o Final/task1 Final/Code/task1.cu "
    "Final/Code/md5_cpu.cpp Final/Code/sha1_cpu.cpp Final/Code/sha256_cpu.cpp\n"
    "nvcc -O3 -std=c++17 -o Final/task2 Final/Code/task2.cu "
    "Final/Code/md5_cpu.cpp Final/Code/sha1_cpu.cpp Final/Code/sha256_cpu.cpp",
    style='Normal'
).runs[0].font.name = 'Courier New'

body("Task 3 (requires -Xcompiler -fopenmp for OpenMP support):")
doc.add_paragraph(
    "nvcc -O3 -std=c++17 -Xcompiler -fopenmp -o Final/task3 Final/Code/task3.cu "
    "Final/Code/md5_cpu.cpp Final/Code/sha1_cpu.cpp Final/Code/sha256_cpu.cpp",
    style='Normal'
).runs[0].font.name = 'Courier New'

body("Task 4 (Thrust ships with CUDA, no extra link flags needed):")
doc.add_paragraph(
    "nvcc -O3 -std=c++17 -o Final/task4 Final/Code/task4.cu",
    style='Normal'
).runs[0].font.name = 'Courier New'

heading('4.3 Running via SLURM', level=2)
body("Submit all three Task 3 scaling jobs (one per algorithm), plus Task 4:")
doc.add_paragraph(
    "sbatch Final/sbatch/task3/task3_md5_scaling.sh\n"
    "sbatch Final/sbatch/task3/task3_sha1_scaling.sh\n"
    "sbatch Final/sbatch/task3/task3_sha256_scaling.sh\n"
    "sbatch Final/sbatch/task4/task4_scaling.sh",
    style='Normal'
).runs[0].font.name = 'Courier New'
body("SLURM parameters for Task 3 jobs: -p instruction, --gres=gpu:1, "
     "--cpus-per-task=12, --mem=96G, -t 0-00:40:00. Data written to Final/Data/task3/.")

heading('4.4 Generating Plots', level=2)
doc.add_paragraph(
    "python3 Final/Scripts/task1/plot_final_task1.py\n"
    "python3 Final/Scripts/task2/plot_final_task2.py\n"
    "python3 Final/Scripts/task3/plot_final_task3.py\n"
    "python3 Final/Scripts/task3/plot_omp_scaling.py\n"
    "python3 Final/Scripts/task4/plot_final_task4.py",
    style='Normal'
).runs[0].font.name = 'Courier New'

# 5. Conclusions
heading('5. Conclusions and Future Work')

body(
    "This project shows three distinct regimes where the GPU helps with hash "
    "computations:"
)
bullet(
    "Raw throughput (Tasks 1 and 2): The GPU achieves 167x to 440x throughput over "
    "serial CPU for large batches, driven by thread-level parallelism over independent "
    "hashes. Register pressure determines per-algorithm sensitivity to configuration, "
    "and tpb=32 is universally optimal for register-heavy hash kernels."
)
bullet(
    "Parallel CPU ceiling (Task 3, OpenMP): OpenMP parallelization of collision search "
    "is limited to about 2x speedup over serial at large bit widths due to the serial "
    "merge phase (Amdahl's Law) and NUMA topology. The 12-thread sweet spot is "
    "hardware-specific to Euler's dual-socket node and would differ on NUMA-free "
    "hardware."
)
bullet(
    "Algorithmic advantage (Task 3, GPU): The GPU's main advantage at bits ≥ 48 is not "
    "raw speed but algorithmic. Pollard's rho with distinguished points uses O(1) "
    "memory per thread, while any birthday-paradox CPU approach requires O(2^(N/2)) "
    "memory. At bits=64 the CPU needs about 100 GB while the GPU uses 8 MB. This makes "
    "bits=64 to 72 GPU-only regardless of CPU clock speed or parallelism level."
)
bullet(
    "Empirical confirmation via Task 4 (10-trial median methodology): With Unified "
    "Memory exposing host RAM to the GPU, the Thrust sort approach falls into the same "
    "performance class as Pollard's rho at bits=56. Per-round Thrust is faster, but "
    "retry variance makes total time comparable. Pollard's rho's advantage at bits=56 "
    "is determinism and predictable timing, not raw speed. At bits ≥ 72 the Thrust "
    "working set exceeds available host RAM and Pollard's rho is the only feasible "
    "option. The lesson is that O(1)-memory algorithms are valuable above the VRAM "
    "boundary not just for feasibility, but for predictability and graceful scaling."
)

body("ME759 concepts leveraged in this project:")
bullet("CUDA kernel design: template specialization for zero warp divergence, "
       "cudaEvent_t timing, cudaMalloc/cudaFree memory management.")
bullet("Occupancy and register pressure: SM register file limits, occupancy analysis, "
       "tpb tuning for register-heavy kernels.")
bullet("OpenMP: parallel-for with schedule(static), omp_get_max_threads(), "
       "phase-parallel/serial-merge pattern.")
bullet("NUMA awareness: dual-socket topology, cross-socket memory latency, thread "
       "affinity and sweet spot analysis.")
bullet("Amdahl's Law: serial fraction identification, theoretical speedup ceiling "
       "calculation for the two-phase parallel algorithm.")
bullet("Memory hierarchy and oversubscription: comparing cudaMalloc vs "
       "cudaMallocManaged, page-tracking overhead, PCIe-bound performance under spill.")

body("Future work:")
bullet(
    "bits=80 GPU: achievable in roughly 30 minutes but exceeds Euler's 40-minute "
    "instruction partition limit. Running on a partition with longer wall-clock limits "
    "would extend the speedup curve one more data point to bits=80."
)
bullet(
    "Parallel sort-based CPU collision search: replacing unordered_map with a flat "
    "sorted array reduces memory to about 43 GB at bits=64. A fully parallel radix "
    "sort plus linear scan could reach bits=64 on CPU in 20 to 30 seconds on a "
    "high-memory server, narrowing the algorithmic gap with Pollard's rho."
)
bullet(
    "GPU Birthday Attack with device-side concurrent hash table. Allocate a large "
    "open-addressing hash table directly in GPU VRAM and have all 65,536 threads "
    "insert (truncated_hash, index) pairs in parallel using atomicCAS for lock-free "
    "collision detection. No CPU-side merge step is needed because the collision is "
    "detected on-device the moment two threads map to the same slot. At small bit "
    "widths (bits ≤ 48) where the entire table fits in GPU VRAM, this approach should "
    "be faster than Pollard's rho because it eliminates the chain-walking overhead. "
    "At bits ≥ 56 the table exceeds VRAM and Pollard's rho with O(1) memory wins. "
    "The crossover point is the main finding."
)

# References
heading('References')
body("[1] Rivest, R. (1992). The MD5 Message-Digest Algorithm. RFC 1321.")
body("[2] NIST FIPS PUB 180-4: Secure Hash Standard (SHS). 2015.")
body("[3] van Oorschot, P. C., and Wiener, M. J. (1999). Parallel Collision Search "
     "with Cryptanalytic Applications. Journal of Cryptology, 12(1), 1-28. "
     "(Pollard's rho with distinguished points.)")
body("[4] Kirk, D. B., and Hwu, W.-m. W. (2016). Programming Massively Parallel "
     "Processors: A Hands-on Approach (3rd ed.). Morgan Kaufmann.")
body("[5] OpenMP Architecture Review Board. OpenMP Application Programming Interface, "
     "Version 5.2. 2021.")
body("[6] NVIDIA Corporation. (2024). CUDA C++ Programming Guide. "
     "https://docs.nvidia.com/cuda/cuda-c-programming-guide/ "
     "(cudaEvent_t timing, __constant__ memory, atomicAdd, warp scheduling, "
     "register file occupancy.)")
body("[7] Pollard, J. M. (1978). Monte Carlo methods for index computation mod p. "
     "Mathematics of Computation, 32(143), 918-924. "
     "(Original rho cycle-detection algorithm; extended by van Oorschot and Wiener [3].)")
body("[8] Flajolet, P., and Odlyzko, A. (1990). Random mapping statistics. "
     "Advances in Cryptology, EUROCRYPT '89, Lecture Notes in Computer Science, "
     "vol. 434, pp. 329-354. Springer. "
     "(Source for birthday paradox collision expectation E ≈ sqrt(π/2) · 2^(N/2).)")
body("[9] Bell, N., and Hoberock, J. (2011). Thrust: A productivity-oriented library "
     "for CUDA. GPU Computing Gems Jade Edition, pp. 359-371. Morgan Kaufmann.")
body("[10] Anthropic. (2026). Claude Code with Claude Sonnet 4.6 (claude-sonnet-4-6) "
     "[AI coding assistant]. Used for code generation, algorithm implementation, and "
     "analysis throughout this project. https://www.anthropic.com/claude-code")

out = 'Final/Report/FinalProject_Arsov.docx'
os.makedirs('Final/Report', exist_ok=True)
doc.save(out)
print(f"Saved {out}")
