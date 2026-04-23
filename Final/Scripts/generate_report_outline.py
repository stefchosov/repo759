"""
Generates Final/Report/FinalProject_Arsov_Outline.docx — concise intermediate submission.
Run from repo root: python3 Final/Scripts/generate_report_outline.py
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def body(text):
    doc.add_paragraph(text)

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(val)
    doc.add_paragraph()

# ── Title ─────────────────────────────────────────────────────────────────────
p = doc.add_heading('ME/CS/ECE 759 Final Project — Intermediate Submission', 0)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('CPU vs GPU Performance and Truncated-Collision Analysis of MD5, SHA-1, SHA-256')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True
p = doc.add_paragraph('Stefan Arsov  |  sarsov@epic.com  |  Spring 2026')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
body('Git repo: https://github.com/stefchosov/repo759')
doc.add_paragraph()

# ── Abstract ──────────────────────────────────────────────────────────────────
h1('Abstract')
body(
    'Benchmarks MD5, SHA-1, and SHA-256 across CPU serial, OpenMP (12-thread), '
    'and GPU (CUDA) for (1) raw throughput vs batch size and tpb configuration, '
    'and (2) truncated collision search over bit widths 16–72. '
    'GPU achieves 167–440× throughput speedup and up to 260× collision-search '
    'speedup at bits=56, and is the only feasible implementation at bits≥64 '
    'using Pollard\'s rho with O(1) memory. Future Task 4 proposes comparing '
    'Pollard\'s rho against a GPU birthday hash table and Thrust parallel sort.'
)

# ── 1. Problem Statement ──────────────────────────────────────────────────────
h1('1. Problem Statement')
body(
    'Cryptographic hash functions underpin authentication, integrity checking, '
    'and digital signatures. Truncated variants (N-bit prefixes) appear in SSH '
    'fingerprints, Git object IDs, and protocol MACs. This project quantifies '
    'how GPU parallelism changes both raw hash throughput and the feasibility of '
    'truncated collision search — directly relevant to practical cryptanalysis scenarios.'
)
bullet('Task 1: GPU vs CPU hash throughput as a function of n (batch size)')
bullet('Task 2: GPU throughput vs threads-per-block (tpb) configuration')
bullet('Task 3: Truncated collision search — CPU serial vs OpenMP vs GPU Pollard\'s rho, bits={16..72}')
bullet('Task 4 (proposed): GPU algorithm comparison — Pollard\'s rho vs GPU birthday table vs Thrust sort')

# ── 2. Solution Description ───────────────────────────────────────────────────
h1('2. Solution Description')

h2('Code Structure  (Final/Code/)')
bullet('task1.cu — n parallel hashes, cudaEvent_t timing, reports MH/s')
bullet('task2.cu — same kernel, sweeps tpb 32→1024')
bullet('task3.cu — CPU serial (unordered_map), OpenMP 2-phase, GPU Pollard\'s rho with DPs; 128-bit extension for bits=72')
bullet('md5_cpu.cpp / sha1_cpu.cpp / sha256_cpu.cpp — portable CPU reference implementations')

h2('Key Design Decisions')
bullet('Template kernels <int ALGO>: zero warp divergence, compiler emits 3 specializations')
bullet('Pollard\'s rho: 65,536 threads each walk independent Markov chain; DP = value with N/4 low bits zero; O(1) memory per thread (8 MB total)')
bullet('OpenMP: Phase 1 parallel hash into hbuf/ibuf, Phase 2 serial merge into unordered_map; 12-thread cap (NUMA sweet spot on dual-socket Euler node)')
bullet('128-bit extension: uint128 struct (lo/hi uint64_t pair), wide hash functions with 16-byte input, pollard_kernel_wide template — enables bits=72')
bullet('CPU_MAX_BITS=56 with --mem=96G; bits=64+ GPU-only (CPU needs ~100 GB for unordered_map at bits=64)')

# ── 3. Results ────────────────────────────────────────────────────────────────
h1('3. Results')

h2('Task 1 — Throughput at n = 10,000,000')
add_table(
    ['Algorithm', 'CPU Serial (MH/s)', 'OpenMP-4 (MH/s)', 'GPU (MH/s)', 'GPU Speedup'],
    [
        ['MD5',     '6.53',  '22.73', '2694.6', '413×'],
        ['SHA-1',   '2.78',   '9.34', '1224.8', '440×'],
        ['SHA-256', '3.94',   '9.54',  '658.9', '167×'],
    ]
)
body('GPU breaks even with CPU serial at n ≈ 200,000 (kernel launch overhead ~2 ms dominates below this).')

h2('Task 2 — tpb Sweep (MD5, n = 10M)')
add_table(
    ['tpb', 'MD5 (MH/s)', 'SHA-1 (MH/s)', 'SHA-256 (MH/s)'],
    [
        ['32 (optimal)', '3021', '1202', '575'],
        ['256 (default)', '2278', '1022', '537'],
        ['1024', '2375', '1003', '557'],
    ]
)
body('tpb=32 optimal for all algorithms. Cause: register pressure limits SM occupancy; larger blocks claim more of the 65K-register file, reducing concurrent warps. MD5 most sensitive (−27%), SHA-256 least (−6%).')

h2('Task 3 — Collision Search (selected results, time in ms)')
add_table(
    ['Algo', 'Bits', 'CPU Serial', 'OMP-12', 'GPU', 'GPU/CPU Speedup'],
    [
        ['MD5',     '48', '4,234',   '2,327',   '544',   '7.8×'],
        ['MD5',     '56', '554,672', '359,864', '2,133', '260×'],
        ['MD5',     '64', '—',       '—',       '8,488', '(CPU infeasible)'],
        ['SHA-1',   '56', '74,417',  '35,765',  '2,375', '31×'],
        ['SHA-256', '56', '316,526', '181,146', '1,930', '164×'],
    ]
)
bullet('OMP speedup over CPU serial: only 1.5–2.1× at bits=56 — serial merge phase (Amdahl) and NUMA dominate')
bullet('12-thread OMP sweet spot confirmed by dedicated scaling experiment (sweep 1–32 threads): NUMA boundary at socket edge, L3 cache saturation, Amdahl ceiling')
bullet('GPU is the only feasible approach at bits≥64 — CPU birthday attack needs ~100 GB for unordered_map; GPU Pollard\'s rho uses 8 MB')
bullet('Birthday paradox validation: cpu_count / expected ≈ 1.0 across all algorithms and bit widths — confirms correct truncated hash extraction')

h2('Task 4 — Proposed GPU Algorithm Comparison')
body('Compare three GPU collision strategies to find crossover points:')
bullet(
    'GPU Birthday Attack (device-side atomicCAS hash table): '
    'fastest at bits≤48 — no chain overhead, collision detected on-device; '
    'limited by VRAM (4 GB table at bits=56 saturates 8 GB GPU).'
)
bullet(
    'Thrust Parallel Sort (thrust::sort_by_key + adjacent scan): '
    'fully GPU-resident, cache-friendly radix sort; '
    'expected fastest at small bits, crossover with Pollard\'s rho ~bits=52–56 as working set fills VRAM.'
)
bullet(
    'Pollard\'s Rho (current, O(1) memory): '
    'dominant at bits≥56 where neither table nor sort fits in VRAM; '
    'only viable approach at bits≥64.'
)
body('Outcome: a complete GPU algorithm selection guide mapping bit width to optimal strategy.')

# ── 4. Deliverables ───────────────────────────────────────────────────────────
h1('4. Deliverables')
bullet('Final/Code/ — task1.cu, task2.cu, task3.cu, CPU hash implementations')
bullet('Final/sbatch/ — SLURM scripts for all tasks (-p instruction, --gres=gpu:1, --mem=96G, 12 CPUs)')
bullet('Final/Scripts/ — Python plot generators for all tasks')
bullet('Final/Data/ — .dat output files and PDF/PNG plots')
body('Compile Task 3: nvcc -O3 -std=c++17 -Xcompiler -fopenmp -o Final/task3 Final/Code/task3.cu Final/Code/md5_cpu.cpp Final/Code/sha1_cpu.cpp Final/Code/sha256_cpu.cpp')
body('Run: OMP_NUM_THREADS=12 ./Final/task3 <algo>  (or via sbatch scripts)')

# ── 5. Conclusions ────────────────────────────────────────────────────────────
h1('5. Conclusions and Future Work')
bullet('GPU throughput advantage (167–440×) comes from embarrassingly parallel independent hashes; tpb=32 is optimal due to register pressure — a free 7–33% gain over the tpb=256 default.')
bullet('OpenMP collision search is bottlenecked by Amdahl\'s serial merge and NUMA; the GPU\'s advantage at large bits is algorithmic (O(1) memory vs O(2^(N/2))), not just faster clocks.')
bullet('Task 4 will quantify the GPU algorithm selection tradeoff: birthday table wins at small bits, Thrust sort in the middle, Pollard\'s rho at large bits.')
bullet('ME759 concepts applied: CUDA template kernels, occupancy/register analysis, OpenMP parallel-for, NUMA topology, Amdahl\'s Law, cudaEvent timing.')

h1('References')
body('[1] Rivest (1992). RFC 1321: MD5 Message-Digest Algorithm.')
body('[2] NIST FIPS 180-4: Secure Hash Standard. 2015.')
body('[3] van Oorschot & Wiener (1999). Parallel Collision Search with Cryptanalytic Applications. J. Cryptology 12(1).')
body('[4] Kirk & Hwu (2016). Programming Massively Parallel Processors, 3rd ed.')
body('[5] Anthropic. (2026). Claude Code with Claude Sonnet 4.6 (claude-sonnet-4-6) '
     '[AI coding assistant]. Used for code generation, algorithm implementation, and '
     'analysis throughout this project. https://www.anthropic.com/claude-code')

out = 'Final/Report/FinalProject_Arsov_Outline.docx'
doc.save(out)
print(f'Saved {out}')
